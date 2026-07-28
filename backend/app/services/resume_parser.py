import io
import re
from datetime import date
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.models.profile import CandidateProfile


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


INVALID_NAME_HEADINGS = {
    "PERSONAL INFORMATION",
    "PERSONAL DETAILS",
    "PERSONAL PROFILE",
    "CURRICULUM VITAE",
    "CURRICULUM VITAE CV",
    "RESUME",
    "CV",
    "PROFILE",
    "PROFESSIONAL PROFILE",
    "PROFESSIONAL SUMMARY",
    "CAREER SUMMARY",
    "SUMMARY",
    "ABOUT",
    "ABOUT ME",
    "CONTACT",
    "CONTACT DETAILS",
    "CAREER OBJECTIVE",
    "OBJECTIVE",
    "EDUCATION",
    "EDUCATIONAL QUALIFICATION",
    "WORK EXPERIENCE",
    "EXPERIENCE",
    "EMPLOYMENT HISTORY",
    "TECHNICAL SKILLS",
    "SKILLS",
    "DECLARATION",
    "AGENTIC AUTOMATION",
    "AGENTIC AUTOMATION CERTIFICATION",
    "CERTIFICATION",
    "PERSONAL STATEMENT",
}


FILENAME_IGNORE_WORDS = {
    "resume",
    "cv",
    "curriculum",
    "vitae",
    "profile",
    "latest",
    "updated",
    "final",
    "new",
    "copy",
    "professional",
    "personal",
    "information",
    "uipath",
    "python",
    "developer",
    "engineer",
    "architect",
    "automation",
    "rpa",
}


KNOWN_SKILLS = [
    "UiPath",
    "UiPath Studio",
    "UiPath Orchestrator",
    "REFramework",
    "Document Understanding",
    "AI Center",
    "Maestro",
    "Python",
    "FastAPI",
    "React",
    "TypeScript",
    "JavaScript",
    "Azure",
    "AWS",
    "SAP",
    "SQL",
    "PostgreSQL",
    "MySQL",
    "Docker",
    "Git",
    "GitHub",
    "Railway",
    "OpenAI",
    "Gemini",
    "LangGraph",
    "Power Automate",
    "Power Platform",
    "Automation Anywhere",
    "Blue Prism",
    "Machine Learning",
    "Artificial Intelligence",
    "Agentic AI",
    "Generative AI",
    "REST API",
    "Microsoft Graph",
    "Orchestrator",
    "RPA",
    "OCR",
    "NLP",
    "Computer Vision",
    "C#",
    "VB.NET",
    "CI/CD",
    "DevOps",
    "ServiceNow",
    "Citrix",
]


MONTH_LOOKUP = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}


def clean_text(value: str | None) -> str:
    """Normalize whitespace and remove unsupported characters."""
    if not value:
        return ""

    cleaned = value.replace("\x00", " ")
    cleaned = cleaned.replace("\u200b", " ")
    cleaned = cleaned.replace("\ufeff", " ")
    cleaned = cleaned.replace("￾", "")

    return re.sub(
        r"[ \t]+",
        " ",
        cleaned,
    ).strip()


def clean_multiline_text(value: str) -> str:
    """Normalize extracted resume text while preserving line breaks."""
    lines = [
        clean_text(line)
        for line in value.splitlines()
    ]

    return "\n".join(
        line
        for line in lines
        if line
    )


def normalize_heading(value: str) -> str:
    """Convert text into a comparable uppercase heading."""
    normalized = re.sub(
        r"[^A-Za-z ]+",
        " ",
        value,
    )

    return clean_text(normalized).upper()


def text_lines(text: str) -> list[str]:
    """Return non-empty normalized resume lines."""
    return [
        clean_text(line)
        for line in text.splitlines()
        if clean_text(line)
    ]


def is_invalid_name(value: str) -> bool:
    """Return True when the value resembles a resume heading."""
    normalized = normalize_heading(value)

    if not normalized:
        return True

    if normalized in INVALID_NAME_HEADINGS:
        return True

    words = normalized.split()

    if len(words) > 5:
        return True

    invalid_phrases = (
        "PERSONAL INFORMATION",
        "PERSONAL DETAILS",
        "PROFESSIONAL SUMMARY",
        "CAREER OBJECTIVE",
        "WORK EXPERIENCE",
        "TECHNICAL SKILLS",
        "EDUCATIONAL QUALIFICATION",
        "AGENTIC AUTOMATION",
        "CERTIFICATION",
        "PERSONAL STATEMENT",
    )

    return any(
        phrase in normalized
        for phrase in invalid_phrases
    )


def extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF resume."""
    reader = PdfReader(io.BytesIO(content))

    pages: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text() or ""

        if page_text:
            pages.append(page_text)

    return clean_multiline_text(
        "\n".join(pages)
    )


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX paragraphs and tables."""
    document = Document(io.BytesIO(content))

    values: list[str] = []

    for paragraph in document.paragraphs:
        value = clean_text(paragraph.text)

        if value:
            values.append(value)

    for table in document.tables:
        for row in table.rows:
            row_values: list[str] = []

            for cell in row.cells:
                value = clean_text(cell.text)

                if value:
                    row_values.append(value)

            if row_values:
                values.append(" | ".join(row_values))

    return clean_multiline_text(
        "\n".join(values)
    )


def extract_text(
    filename: str,
    content: bytes,
) -> str:
    """Extract text based on the uploaded resume extension."""
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported resume format. "
            "Upload a PDF, DOCX, or TXT file."
        )

    if extension == ".pdf":
        text = extract_pdf_text(content)

    elif extension == ".docx":
        text = extract_docx_text(content)

    else:
        text = content.decode(
            "utf-8",
            errors="ignore",
        )

        text = clean_multiline_text(text)

    if not clean_text(text):
        raise ValueError(
            "No readable text was found in the resume."
        )

    return text


def first_match(
    pattern: str,
    text: str,
    flags: int = re.IGNORECASE,
) -> str:
    """Return the first captured value from a regex match."""
    match = re.search(
        pattern,
        text,
        flags,
    )

    if not match:
        return ""

    return clean_text(match.group(1))


def remove_trailing_punctuation(value: str) -> str:
    return value.rstrip(
        ".,;:|)]}>"
    )


def extract_email(text: str) -> str:
    value = first_match(
        r"\b([A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,})\b",
        text,
    )

    return remove_trailing_punctuation(value)


def extract_phone(text: str) -> str:
    candidates = re.findall(
        r"(?<!\d)"
        r"(\+?\s*\(?\d{1,3}\)?[\d\s().\-]{7,}\d)"
        r"(?!\d)",
        text,
    )

    for candidate in candidates:
        cleaned = clean_text(candidate)

        digits = re.sub(
            r"\D",
            "",
            cleaned,
        )

        if 10 <= len(digits) <= 15:
            return cleaned

    return ""


def extract_linkedin(text: str) -> str:
    value = first_match(
        r"((?:https?://)?"
        r"(?:www\.)?"
        r"linkedin\.com/in/"
        r"[A-Za-z0-9_%\-]+"
        r"/?)",
        text,
    )

    return remove_trailing_punctuation(value)


def extract_github(text: str) -> str:
    value = first_match(
        r"((?:https?://)?"
        r"(?:www\.)?"
        r"github\.com/"
        r"[A-Za-z0-9_.\-]+"
        r"/?)",
        text,
    )

    return remove_trailing_punctuation(value)


def extract_portfolio(text: str) -> str:
    urls = re.findall(
        r"(https?://[^\s<>()|]+)",
        text,
        re.IGNORECASE,
    )

    for url in urls:
        lower_url = url.lower()

        if (
            "linkedin.com" not in lower_url
            and "github.com" not in lower_url
            and "credentials.uipath.com" not in lower_url
        ):
            return remove_trailing_punctuation(url)

    return ""


def candidate_name_from_filename(
    filename: str,
) -> str:
    """Derive a safe candidate name from the filename."""
    stem = Path(filename).stem

    stem = re.sub(
        r"\(\d+\)$",
        "",
        stem,
    )

    stem = re.sub(
        r"[\[\]{}()]",
        " ",
        stem,
    )

    tokens = re.split(
        r"[_\-\s]+",
        stem,
    )

    usable_tokens: list[str] = []

    for token in tokens:
        cleaned = re.sub(
            r"[^A-Za-z.' ]",
            "",
            token,
        ).strip()

        if not cleaned:
            continue

        if cleaned.lower() in FILENAME_IGNORE_WORDS:
            continue

        if cleaned.isdigit():
            continue

        usable_tokens.append(cleaned)

    if not usable_tokens:
        return ""

    return " ".join(
        word.capitalize()
        for word in usable_tokens[:3]
    )


def looks_like_person_name(value: str) -> bool:
    """Validate whether a line could be a person's name."""
    cleaned = clean_text(value)

    if not cleaned:
        return False

    if is_invalid_name(cleaned):
        return False

    lowered = cleaned.lower()

    invalid_markers = (
        "@",
        "http://",
        "https://",
        "linkedin",
        "github",
        "phone",
        "mobile",
        "email",
        "address",
        "summary",
        "experience",
        "objective",
        "skills",
        "education",
        "certification",
        "automation",
        "developer",
        "engineer",
        "architect",
        "consultant",
        "manager",
        "lead",
        "profile",
        "statement",
    )

    if any(
        marker in lowered
        for marker in invalid_markers
    ):
        return False

    if ":" in cleaned:
        return False

    words = cleaned.split()

    if not 2 <= len(words) <= 4:
        return False

    for word in words:
        if not re.fullmatch(
            r"[A-Za-z][A-Za-z.'\-]*",
            word,
        ):
            return False

    return True


def extract_name_after_personal_information(
    text: str,
) -> str:
    """
    Extract the candidate name near the PERSONAL INFORMATION section.

    Handles layouts where sidebar headings appear before the actual name.
    """
    lines = text_lines(text)

    heading_index = -1

    for index, line in enumerate(lines[:30]):
        if (
            normalize_heading(line)
            == "PERSONAL INFORMATION"
        ):
            heading_index = index
            break

    if heading_index >= 0:
        nearby_lines = lines[
            heading_index + 1:
            heading_index + 12
        ]

        for line in nearby_lines:
            if looks_like_person_name(line):
                return line

    return ""


def extract_name_from_contact_details(
    text: str,
) -> str:
    """
    Search for a name close to address, mobile, email, or LinkedIn lines.
    """
    lines = text_lines(text)

    for index, line in enumerate(lines[:40]):
        lowered = line.lower()

        if not any(
            marker in lowered
            for marker in (
                "address:",
                "mobile:",
                "email:",
                "linkedin profile:",
            )
        ):
            continue

        for previous_line in reversed(
            lines[max(0, index - 5):index]
        ):
            if looks_like_person_name(
                previous_line
            ):
                return previous_line

    return ""


def extract_name_from_text(text: str) -> str:
    """Extract the most likely candidate name."""
    personal_information_name = (
        extract_name_after_personal_information(
            text
        )
    )

    if personal_information_name:
        return personal_information_name

    contact_name = extract_name_from_contact_details(
        text
    )

    if contact_name:
        return contact_name

    for line in text_lines(text)[:40]:
        if looks_like_person_name(line):
            return line

    return ""


def split_full_name(
    full_name: str,
) -> tuple[str, str, str]:
    cleaned = clean_text(full_name)

    if not cleaned or is_invalid_name(cleaned):
        return "", "", ""

    words = cleaned.split()

    if len(words) == 1:
        return (
            words[0],
            "",
            cleaned,
        )

    return (
        words[0],
        " ".join(words[1:]),
        cleaned,
    )


def extract_name(
    text: str,
    filename: str,
) -> tuple[str, str, str]:
    """Resolve the candidate name from resume text or filename."""
    extracted_name = extract_name_from_text(text)

    if extracted_name:
        result = split_full_name(
            extracted_name
        )

        if result[2]:
            return result

    filename_name = candidate_name_from_filename(
        filename
    )

    if filename_name:
        return split_full_name(
            filename_name
        )

    return "", "", ""


def extract_labeled_value(
    labels: list[str],
    text: str,
) -> str:
    """Extract a value appearing after a known label."""
    label_pattern = "|".join(
        re.escape(label)
        for label in labels
    )

    pattern = (
        rf"(?:{label_pattern})"
        rf"\s*[:\-]\s*"
        rf"([^\n\r|]+)"
    )

    return first_match(pattern, text)


def clean_profile_value(value: str) -> str:
    """Clean extracted profile values."""
    cleaned = clean_text(value)

    cleaned = cleaned.strip(
        " -–—:;,|"
    )

    if not cleaned:
        return ""

    normalized = normalize_heading(cleaned)

    invalid_values = {
        "WORK EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "EMPLOYMENT HISTORY",
        "CURRENT COMPANY",
        "CURRENT EMPLOYER",
        "CURRENT DESIGNATION",
        "CURRENT ROLE",
        "JOB TITLE",
        "DESIGNATION",
        "EXTENSIVE EXPERIENCE",
    }

    if normalized in invalid_values:
        return ""

    return cleaned


def extract_city(text: str) -> str:
    labeled_value = extract_labeled_value(
        [
            "Current City",
            "Current Location",
            "City",
            "Location",
        ],
        text,
    )

    return clean_profile_value(
        labeled_value
    )


def extract_country(text: str) -> str:
    labeled_value = extract_labeled_value(
        [
            "Country",
            "Nationality",
        ],
        text,
    )

    return clean_profile_value(
        labeled_value
    )


def normalize_company_name(value: str) -> str:
    """Clean company-name formatting."""
    cleaned = clean_profile_value(value)

    cleaned = re.split(
        r"\s+[•]",
        cleaned,
        maxsplit=1,
    )[0]

    cleaned = re.sub(
        r"\s{2,}",
        " ",
        cleaned,
    )

    return cleaned.strip(
        " -–—:;,|"
    )


def extract_current_role_and_company(
    text: str,
) -> tuple[str, str]:
    """
    Extract the most recent designation and company.

    Supports patterns such as:
    April 2023 - Till Date RPA Sr. Automation Developer at Zelis India Pvt Ltd
    """
    normalized_text = clean_multiline_text(text)

    patterns = [
        (
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec|"
            r"January|February|March|April|June|July|August|September|"
            r"October|November|December)"
            r"\s+\d{4}"
            r"\s*[-–—]\s*"
            r"(?:Till\s+Date|Present|Current)"
            r"\s+"
            r"(.+?)"
            r"\s+at\s+"
            r"([^\n\r]+)"
        ),
        (
            r"(?:Till\s+Date|Present|Current)"
            r"\s+"
            r"(.+?)"
            r"\s+at\s+"
            r"([^\n\r]+)"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            normalized_text,
            re.IGNORECASE,
        )

        if not match:
            continue

        designation = clean_profile_value(
            match.group(1)
        )

        company = normalize_company_name(
            match.group(2)
        )

        company = re.split(
            r"\s+(?:Strategic|Liaising|Responsibilities|Role|Project)\b",
            company,
            maxsplit=1,
            flags=re.IGNORECASE,
        )[0]

        company = re.sub(
            r"^(?:RPA)?\s*",
            "",
            company,
            flags=re.IGNORECASE,
        )

        if designation and company:
            return designation, company

    return "", ""


def extract_current_company(text: str) -> str:
    """Extract the current employer."""
    _, company = extract_current_role_and_company(
        text
    )

    if company:
        return company

    labeled_value = extract_labeled_value(
        [
            "Current Company",
            "Current Employer",
            "Present Company",
            "Present Employer",
            "Organization",
            "Employer",
            "Company",
        ],
        text,
    )

    labeled_value = normalize_company_name(
        labeled_value
    )

    invalid_company_phrases = (
        "extensive experience",
        "professional experience",
        "work experience",
    )

    if any(
        phrase in labeled_value.lower()
        for phrase in invalid_company_phrases
    ):
        return ""

    return labeled_value


def extract_designation(text: str) -> str:
    """Extract the current or most recent designation."""
    designation, _ = (
        extract_current_role_and_company(
            text
        )
    )

    if designation:
        designation = re.sub(
            r"^RPA(?=Sr\.|Senior|Developer|Automation)",
            "RPA ",
            designation,
            flags=re.IGNORECASE,
        )

        designation = re.sub(
            r"(?i)\bRPASr\.",
            "RPA Sr.",
            designation,
        )

        return clean_profile_value(
            designation
        )

    labeled_value = extract_labeled_value(
        [
            "Current Designation",
            "Current Role",
            "Current Job Title",
            "Designation",
            "Job Title",
            "Role",
            "Position",
        ],
        text,
    )

    return clean_profile_value(
        labeled_value
    )


def parse_month_year(
    month_name: str,
    year_value: str,
) -> tuple[int, int] | None:
    month = MONTH_LOOKUP.get(
        month_name.lower()
    )

    if not month:
        return None

    try:
        year = int(year_value)
    except ValueError:
        return None

    return year, month


def extract_employment_start_dates(
    text: str,
) -> list[tuple[int, int]]:
    """Extract employment start month/year values."""
    pattern = (
        r"\b("
        r"Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|"
        r"May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|"
        r"Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?"
        r")"
        r"\s+(\d{4})"
        r"\s*[-–—]"
    )

    dates: list[tuple[int, int]] = []

    for month_name, year_value in re.findall(
        pattern,
        text,
        re.IGNORECASE,
    ):
        parsed = parse_month_year(
            month_name,
            year_value,
        )

        if parsed:
            dates.append(parsed)

    return dates


def calculate_experience_from_dates(
    text: str,
) -> str:
    """
    Calculate experience from the earliest employment start date.

    Returns years with one decimal when needed.
    """
    start_dates = extract_employment_start_dates(
        text
    )

    if not start_dates:
        return ""

    earliest_year, earliest_month = min(
        start_dates
    )

    today = date.today()

    total_months = (
        (today.year - earliest_year) * 12
        + today.month
        - earliest_month
    )

    if total_months <= 0:
        return ""

    years = total_months / 12

    rounded_years = round(
        years,
        1,
    )

    if rounded_years.is_integer():
        return str(int(rounded_years))

    return f"{rounded_years:.1f}"


def extract_experience_years(text: str) -> str:
    """Extract total professional experience."""
    explicit_patterns = [
        (
            r"(?:total\s+)?(?:professional\s+)?experience"
            r"\s*[:\-]?\s*"
            r"(\d{1,2}(?:\.\d+)?)\+?"
            r"\s*(?:years?|yrs?)"
        ),
        (
            r"(\d{1,2}(?:\.\d+)?)\+?"
            r"\s*(?:years?|yrs?)"
            r"\s+(?:of\s+)?"
            r"(?:total\s+)?"
            r"(?:professional\s+)?experience"
        ),
        (
            r"(?:over|more\s+than|around|approximately)"
            r"\s+(\d{1,2}(?:\.\d+)?)\+?"
            r"\s*(?:years?|yrs?)"
            r"(?:\s+of)?\s+experience"
        ),
    ]

    for pattern in explicit_patterns:
        value = first_match(
            pattern,
            text,
        )

        if value:
            return value

    return calculate_experience_from_dates(
        text
    )


def extract_notice_period(text: str) -> str:
    return clean_profile_value(
        extract_labeled_value(
            [
                "Notice Period",
                "Availability",
                "Available From",
            ],
            text,
        )
    )


def extract_current_salary(text: str) -> str:
    return clean_profile_value(
        extract_labeled_value(
            [
                "Current Salary",
                "Current CTC",
                "Current Compensation",
            ],
            text,
        )
    )


def extract_expected_salary(text: str) -> str:
    return clean_profile_value(
        extract_labeled_value(
            [
                "Expected Salary",
                "Expected CTC",
                "Expected Compensation",
            ],
            text,
        )
    )


def extract_skills(text: str) -> str:
    """Extract known skills without duplicates."""
    found: list[str] = []

    for skill in KNOWN_SKILLS:
        if re.search(
            rf"(?<![A-Za-z0-9])"
            rf"{re.escape(skill)}"
            rf"(?![A-Za-z0-9])",
            text,
            re.IGNORECASE,
        ):
            if skill not in found:
                found.append(skill)

    return ", ".join(found)


def parse_resume(
    filename: str,
    content: bytes,
) -> CandidateProfile:
    """
    Parse an uploaded resume into a CandidateProfile.

    Missing values remain empty. No data is invented.
    """
    text = extract_text(
        filename=filename,
        content=content,
    )

    first_name, last_name, full_name = extract_name(
        text=text,
        filename=filename,
    )

    return CandidateProfile(
        firstName=first_name,
        lastName=last_name,
        fullName=full_name,
        email=extract_email(text),
        phone=extract_phone(text),
        city=extract_city(text),
        country=extract_country(text),
        linkedin=extract_linkedin(text),
        github=extract_github(text),
        portfolio=extract_portfolio(text),
        currentCompany=extract_current_company(
            text
        ),
        designation=extract_designation(
            text
        ),
        experienceYears=extract_experience_years(
            text
        ),
        noticePeriod=extract_notice_period(
            text
        ),
        currentSalary=extract_current_salary(
            text
        ),
        expectedSalary=extract_expected_salary(
            text
        ),
        skills=extract_skills(text),
        coverLetter="",
        resume=filename,
    )
